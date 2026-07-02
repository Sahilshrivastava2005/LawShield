"""
parser.py
Low-level document parsers for PDF, DOCX, and plain-text files.

Strategy (in priority order for PDFs):
  1. PyMuPDF (fitz)    — fast, handles most digitally-created PDFs.
  2. pdfplumber        — better table/column layout extraction.
  3. OCR via ocr.py    — last resort for fully scanned (image-only) PDFs.

Each parser returns a list of :class:`PageResult` dicts so callers always
know *which page* each piece of text came from — useful for citation and
chunk metadata.
"""

from __future__ import annotations

import logging
from dataclasses import dataclass, field
from typing import List, Optional

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Data model
# ---------------------------------------------------------------------------

@dataclass
class PageResult:
    """Text and metadata extracted from a single document page."""
    page_number: int          # 1-indexed
    text: str                 # Raw extracted text for this page
    tables: List[str] = field(default_factory=list)   # Stringified tables (pdfplumber)
    is_scanned: bool = False  # True when the page had no selectable text (OCR used)
    word_count: int = 0

    def __post_init__(self) -> None:
        self.word_count = len(self.text.split())


def pages_to_text(pages: List[PageResult]) -> str:
    """Concatenate all page texts with a page-break marker."""
    parts = []
    for p in pages:
        parts.append(p.text)
        if p.tables:
            parts.append("\n".join(p.tables))
    return "\n\n".join(filter(None, parts))


# ---------------------------------------------------------------------------
# Scanned-page detection
# ---------------------------------------------------------------------------

_SCANNED_TEXT_THRESHOLD = 20   # fewer than this many chars → treat as scanned


def _is_scanned_page(text: str) -> bool:
    """Return True if the page has so little text it is likely a scanned image."""
    return len(text.strip()) < _SCANNED_TEXT_THRESHOLD


# ---------------------------------------------------------------------------
# PDF — PyMuPDF (fitz)
# ---------------------------------------------------------------------------

def parse_pdf_pymupdf(file_path: str) -> List[PageResult]:
    """
    Fast, reliable extraction using PyMuPDF.

    - Extracts text per page (layout-preserving via "text" mode).
    - Detects scanned pages (image-only) by comparing text length to threshold.
    """
    pages: List[PageResult] = []
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(file_path)
        for i, page in enumerate(doc, start=1):
            # "blocks" mode preserves reading order better than raw get_text()
            text = page.get_text("text")
            scanned = _is_scanned_page(text)
            pages.append(PageResult(page_number=i, text=text, is_scanned=scanned))
        doc.close()
    except Exception as exc:
        logger.error("PyMuPDF parsing failed for %s: %s", file_path, exc)
    return pages


# ---------------------------------------------------------------------------
# PDF — pdfplumber  (table-heavy / complex layout)
# ---------------------------------------------------------------------------

def parse_pdf_pdfplumber(file_path: str) -> List[PageResult]:
    """
    Layout-aware extraction using pdfplumber.

    Also extracts tables and converts them to a CSV-like string so they
    are preserved for downstream processing and chunking.
    """
    pages: List[PageResult] = []
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""
                scanned = _is_scanned_page(text)

                # Extract tables if present
                table_strs: List[str] = []
                try:
                    tables = page.extract_tables()
                    for table in (tables or []):
                        rows = []
                        for row in table:
                            cells = [str(c).strip() if c is not None else "" for c in row]
                            rows.append(" | ".join(cells))
                        table_strs.append("\n".join(rows))
                except Exception as tbl_exc:
                    logger.debug("Table extraction failed page %d: %s", i, tbl_exc)

                pages.append(PageResult(
                    page_number=i,
                    text=text,
                    tables=table_strs,
                    is_scanned=scanned,
                ))
    except Exception as exc:
        logger.error("pdfplumber parsing failed for %s: %s", file_path, exc)
    return pages


# ---------------------------------------------------------------------------
# PDF — OCR fallback for scanned pages
# ---------------------------------------------------------------------------

def parse_pdf_with_ocr(file_path: str) -> List[PageResult]:
    """
    Convert each PDF page to an image and run OCR on it.
    Used when a page has no selectable text (fully scanned document).
    Requires PyMuPDF for rendering and PaddleOCR for recognition.
    """
    pages: List[PageResult] = []
    try:
        import fitz
        from .ocr import extract_text_with_ocr
        import tempfile, os

        doc = fitz.open(file_path)
        for i, page in enumerate(doc, start=1):
            # Render page to image at 200 DPI
            mat = fitz.Matrix(200 / 72, 200 / 72)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp_img:
                tmp_path = tmp_img.name
                pix.save(tmp_path)
            try:
                ocr_text = extract_text_with_ocr(tmp_path)
            finally:
                if os.path.exists(tmp_path):
                    os.remove(tmp_path)
            pages.append(PageResult(page_number=i, text=ocr_text, is_scanned=True))
        doc.close()
    except Exception as exc:
        logger.error("OCR parsing failed for %s: %s", file_path, exc)
    return pages


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def parse_docx_unstructured(file_path: str) -> List[PageResult]:
    """
    Parse DOCX files using the Unstructured library.
    Falls back to python-docx on failure.
    Returns a single 'page' (DOCX has no fixed page concept).
    """
    text = ""
    try:
        from unstructured.partition.docx import partition_docx
        elements = partition_docx(filename=file_path)
        text = "\n".join(str(el) for el in elements if str(el).strip())
    except Exception as exc:
        logger.warning("Unstructured DOCX failed (%s), trying python-docx…", exc)
        text = _parse_docx_fallback(file_path)

    return [PageResult(page_number=1, text=text)]


def _parse_docx_fallback(file_path: str) -> str:
    """Minimal python-docx fallback."""
    try:
        from docx import Document  # type: ignore
        doc = Document(file_path)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception as exc:
        logger.error("python-docx fallback also failed: %s", exc)
        return ""


# ---------------------------------------------------------------------------
# Plain text
# ---------------------------------------------------------------------------

def parse_txt(file_path: str, encoding: str = "utf-8") -> List[PageResult]:
    """
    Read a plain-text file.
    Tries UTF-8 first, then falls back to latin-1.
    """
    for enc in (encoding, "latin-1"):
        try:
            with open(file_path, encoding=enc) as fh:
                text = fh.read()
            return [PageResult(page_number=1, text=text)]
        except UnicodeDecodeError:
            continue
        except Exception as exc:
            logger.error("TXT parsing failed: %s", exc)
            break
    return [PageResult(page_number=1, text="")]
